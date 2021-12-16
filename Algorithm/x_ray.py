from converting_to_docx import *
from Algorithm.export_xray_reports import *
from reports_segmentation import create_new_document

from Keys.global_keys import *

from docx import Document
import docx
import os


class Xray:
    def __init__(self, path: str):
        self.path_xray = path

    def convert_xray_files(self):
        for dir_path, dir_names, filenames in os.walk(self.path_xray):
            for dir_name in dir_names:
                if dir_name.endswith("-"):
                    docx_(root, dir_path + '/' + dir_name, dir_path + '/' + "converted_" + dir_name)

    def prepare_xray_conclusion(self):
        for dir_path, dir_names, filenames in os.walk(self.path_xray):
            for filename in filenames:
                if '.docx' in filename and dir_path.endswith('-'):  #and 'converted_' in dir_path
                    path_new = dir_path.replace('converted_', '')\
                        .lstrip('/home/lg/Dropbox/Conclusion/MEDICAL REPORTS/')\
                        .rstrip('-')\
                        .replace('-/', '/').replace('/', '//')
                    self.segment_else(filename, path_new, dir_path)

    def export_to_database(self):
        for dir_path, dir_names, filenames in os.walk(self.path_xray):
            for filename in filenames:
                if '.docx' in filename and dir_path.endswith('-') \
                        and 'converted_' in dir_path:
                    dirname = dir_path.split('converted_')
                    exp = AutoExport(export_link, root, dir_path, filename, dirname[1].rstrip('-'))
                    exp.export()

    @staticmethod
    def segment_else(prepared_document: str, anatomy_structure_path, dir_path):
        """
        Выявляем патологии не попавшие в список с ключевыми словами
        :return:
        """
        document_path = dir_path + '/' + prepared_document
        try:
            conclusion = Document(document_path)
            new_doc_step1 = docx.Document()
            number_paragraphs = len(conclusion.paragraphs)
            text = f"{anatomy_structure_path}\n{prepared_document.rstrip('.docx')}\n-report-text-below-"
            new_doc_step1.add_paragraph(text)
            create_new_document(new_doc_step1, number_paragraphs, conclusion, dir_path, prepared_document)
        except ValueError:
            print(f"Необработанный документ: {prepared_document}")
        else:
            pass


if __name__ == '__main__':
    xray = Xray('/home/lg/Dropbox/Conclusion/MEDICAL REPORTS/Рентген/')
    # xray.convert_xray_files()
    # xray.prepare_xray_conclusion()
    xray.export_to_database()
