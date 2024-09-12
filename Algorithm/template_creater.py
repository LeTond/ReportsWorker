"""
Шаблонизатор.
Преобразует подготовленный шаблон к заливке в БД
"""
import os
from docx import Document


ROOT_DIR = "root_directory"
EMPTY_STR = ""


class Template:
    def __init__(self):
        self.list_glob = []
        self.path_to_directory = "/path_to_data"

    def template(self):
        """
        Поиск параграфов и вызов соответствующей функции для форматирования
        :return: None
        """
        for fltr in os.listdir(self.path_to_directory):
            path = self.path_to_directory + '/' + fltr
            try:
                if ".docx" in fltr:
                    conclusion = Document(path)
                    for ind, text_ in enumerate(conclusion.paragraphs):
                        if "<P>" in text_.text:
                            self.paragraph(ind, conclusion)
                        if "<S" in text_.text:
                            self.sentence(ind, conclusion)
                        if "<RB>" in text_.text:
                            self.radio_button(ind, conclusion)
                        if "<CHECK>" in text_.text:
                            self.check_box(ind, conclusion)
                        if "<DROP>" in text_.text:
                            self.drop_down(ind, conclusion)
                        if "<INPUT>" in text_.text:
                            self.input_text(ind, conclusion)
                        if "<STATIC>" in text_.text:
                            self.static_text(ind, conclusion)
                        if "<RECOMMEND>" in text_.text:
                            self.recommendation(ind, conclusion)
                        if "<CONCLUSIONTEXT>" in text_.text:
                            self.conclusion_text(ind, conclusion)
                    self.save_to_word(path)
            except ValueError:
                print(f"Необработанный документ: {fltr}")

    def paragraph(self, ind: int, conclusion: any):
        """
        Создание формы нового параграфа для шаблона
        :param ind: Номер строки в тексте
        :param conclusion: текст шаблона
        :return: None
        """
        part = conclusion.paragraphs[ind + 1].text
        self.list_glob.append(f"\nparagraph: {part}")
        if "<COMMENT>" in conclusion.paragraphs[ind].text:
            self.list_glob.append("paragraph_is_comment: true")

    def sentence(self, ind: int, conclusion: any):
        """
        Создание формы нового предложения в параграфе
        :param ind: Номер строки в тексте
        :param conclusion: текст шаблона
        :return: None
        """
        if "<SR>" in conclusion.paragraphs[ind].text:
            self.list_glob.append(f"\n\tsentence: {conclusion.paragraphs[ind + 1].text}\n"
                                  f"\tsentence_type: regular")
        if "<SM>" in conclusion.paragraphs[ind].text:
            self.list_glob.append(f"\n\tsentence: {conclusion.paragraphs[ind + 1].text}\n"
                                  f"\tsentence_type: multiple")
        if "<NEWLINE>" in conclusion.paragraphs[ind].text:
            self.list_glob.append("\tsentence_new_line: true")

    def start_end_text(self, ind: int, conclusion: any) -> int:
        """
        Создание формы начала и конца предложения в параграфе
        :param ind: Номер строки в тексте
        :param conclusion: текст шаблона
        :return: None
        """
        counter = 1
        if "<START " in conclusion.paragraphs[ind + counter].text:
            start = conclusion.paragraphs[ind + counter].text.lstrip("<START ").rstrip("> ")
            self.list_glob.append(f"\t\tsubsentence_enclosing_start_text: {start}")
            counter += 1

        if "<END " in conclusion.paragraphs[ind + counter].text:
            end = conclusion.paragraphs[ind + counter].text.lstrip("<END ").rstrip("> ")
            self.list_glob.append(f"\t\tsubsentence_enclosing_end_text: {end}")
            counter += 1

        if "<STARTSINGLE" in conclusion.paragraphs[ind + counter].text:
            start_single = conclusion.paragraphs[ind + counter].text.lstrip("<STARTSINGLE ").rstrip("> ")
            self.list_glob.append(f"\t\tsubsentence_enclosing_start_text: {start_single}")
            counter += 1

        if "<STARTMULTIPLE" in conclusion.paragraphs[ind + counter].text:
            start_mult = conclusion.paragraphs[ind + counter].text.lstrip("<STARTMULTIPLE ").rstrip("> ")
            self.list_glob.append(f"\t\tsubsentence_enclosing_start_text_plural: {start_mult}")
            counter += 1

        if "<ENDSINGLE" in conclusion.paragraphs[ind + counter].text:
            end_single = conclusion.paragraphs[ind + counter].text.lstrip("<ENDSINGLE ").rstrip("> ")
            self.list_glob.append(f"\t\tsubsentence_enclosing_end_text: {end_single}")
            counter += 1

        if "<ENDMULTIPLE" in conclusion.paragraphs[ind + counter].text:
            end_mult = conclusion.paragraphs[ind + counter].text.lstrip("<ENDMULTIPLE ").rstrip("> ")
            self.list_glob.append(f"\t\tsubsentence_enclosing_end_text_plural: {end_mult}")
            counter += 1

        counter -= 1
        return counter

    def check_box(self, ind: int, conclusion: any):
        """
        Создание формы check_box
        :param ind: Номер строки в тексте
        :param conclusion: текст шаблона
        :return: None
        """
        counter = 1
        add_ind = 0
        self.list_glob.append("\n\t\tsubsentence_type: checkbox\n"
                              "\t\tsubsentence_title:")
        while "<" not in conclusion.paragraphs[ind + counter].text:
            counter += 1
            add_ind += 1
        counter += self.start_end_text(ind + add_ind, conclusion)
        counter = 1
        while "<" not in conclusion.paragraphs[ind + counter].text:
            try:
                part = conclusion.paragraphs[ind + counter].text.split(' [')
                if len(part) == 2:
                    part_1 = part[1].rstrip("] // ")
                    self.list_glob.append(f"\t\t\tcheckbox_ui_title: {part[0]}\n"
                                          f"\t\t\tcheckbox_selected_text: {part_1}"
                                          )
                elif len(part) == 3:
                    part_1 = part[1].strip("] ")
                    part_2 = part[2].rstrip("] // ")
                    self.list_glob.append(f"\t\t\tcheckbox_ui_title: {part[0]}\n"
                                          f"\t\t\tcheckbox_selected_text: {part_1}\n"
                                          f"\t\t\tcheckbox_link_id: {part_2}"
                                          )
                else:
                    print(f"Ошибка check_box в строке: {part}, номер строки: {ind}")

            except IndexError:
                print(f"Ошибка check_box в предложении: {conclusion.paragraphs[ind + counter].text}")
            counter += 1

    def radio_button(self, ind, conclusion):
        """
        Создание формы radio_button
        :param ind: Номер строки в тексте
        :param conclusion: текст шаблона
        :return: None
        """
        counter = 1
        add_ind = 0
        self.list_glob.append("\n\t\tsubsentence_type: radiobutton\n"
                              "\t\tsubsentence_title:")
        while "<" not in conclusion.paragraphs[ind + counter].text:
            counter += 1
            add_ind += 1
        counter += self.start_end_text(ind + add_ind, conclusion)
        counter = 1
        while "<" not in conclusion.paragraphs[ind + counter].text:
            try:
                part = conclusion.paragraphs[ind + counter].text.split(' [')
                if len(part) == 2:
                    part_1 = part[1].rstrip("] //")
                    self.list_glob.append(f"\t\t\tradiobutton_ui_title: {part[0]}\n"
                                          f"\t\t\tradiobutton_selected_text: {part_1}"
                                          )
                elif len(part) == 3:
                    part_1 = part[1].strip("] ")
                    part_2 = part[2].rstrip("] //")
                    self.list_glob.append(f"\t\t\tradiobutton_ui_title: {part[0]}\n"
                                          f"\t\t\tradiobutton_selected_text: {part_1}\n"
                                          f"\t\t\tradiobutton_link_id: {part_2}"
                                          )
                else:
                    print(f"Ошибка radiobutton в строке: {part}, номер строки: {ind}")
            except IndexError:
                print(f"Ошибка radiobutton в предложении: {conclusion.paragraphs[ind + counter].text}")
            counter += 1

    def drop_down(self, ind: int, conclusion: any):
        """
        Создание формы drop_down
        :param ind: Номер строки в тексте
        :param conclusion: текст шаблона
        :return: None
        """
        counter = 1
        add_ind = 0
        self.list_glob.append("\n\t\tsubsentence_type: dropdown\n"
                              "\t\tsubsentence_title:")
        while "<" not in conclusion.paragraphs[ind + counter].text:
            counter += 1
            add_ind += 1
        counter += self.start_end_text(ind + add_ind, conclusion)
        counter = 1
        while "<" not in conclusion.paragraphs[ind + counter].text:
            try:

                part = conclusion.paragraphs[ind + counter].text.split(' [')

                if len(part) == 2:
                    part_1 = part[1].rstrip("] //")
                    self.list_glob.append(f"\t\t\tdropdown_ui_title: {part[0]}\n"
                                          f"\t\t\tdropdown_selected_text: {part_1}"
                                          )
                elif len(part) == 3:
                    part_1 = part[1].strip("] ")
                    part_2 = part[2].rstrip("] //")
                    self.list_glob.append(f"\t\t\tdropdown_ui_title: {part[0]}\n"
                                          f"\t\t\tdropdown_selected_text: {part_1}\n"
                                          f"\t\t\tdropdown_link_id: {part_2}"
                                          )
                else:
                    print(f"Ошибка dropdown в строке: {part}, номер строки: {ind}")

            except IndexError:
                print(f"Ошибка dropdown в предложении: {conclusion.paragraphs[ind + counter].text}")
            counter += 1

    def input_text(self, ind: int, conclusion: any):
        """
        Создание формы input_text
        :param ind: Номер строки в тексте
        :param conclusion: текст шаблона
        :return: None
        """
        counter = 1
        while "<" not in conclusion.paragraphs[ind + counter].text:
            self.list_glob.append("\n\t\tsubsentence_type: input_text\n"
                                  "\t\tsubsentence_title:")
            try:
                part = conclusion.paragraphs[ind + counter].text.split(' [')
                ui_stt = part[0].split("…")
                stt = part[1].rstrip("] //").split("…")
                self.list_glob.append(
                    f"\t\t\tsubsentence_enclosing_input_ui_start_text: {ui_stt[0]}\n"
                    f"\t\t\tsubsentence_enclosing_input_ui_end_text: {ui_stt[1]}\n"
                    f"\t\t\tsubsentence_enclosing_input_start_text: {stt[0]}\n"
                    f"\t\t\tsubsentence_enclosing_input_end_text: {stt[1]}"
                )
            except IndexError:
                print(f"Ошибка input_text в предложении: {conclusion.paragraphs[ind + counter].text}"
                      f", номер строки: {ind}")
            counter += 1

    def static_text(self, ind: int, conclusion: any):
        """
        Создание формы static_text
        :param ind: Номер строки в тексте
        :param conclusion: текст шаблона
        :return: None
        """
        part = conclusion.paragraphs[ind + 1].text
        part_1 = part.lstrip("[").rstrip("] ")
        self.list_glob.append(f"\n\t\tsubsentence_type: static_text"
                              f"\n\t\tsubsentence_title:\n"
                              f"\n\t\t\tstatic_text_value: {part_1}"
                              )

    def recommendation(self, ind: int, conclusion: any):
        """
        paragraph: Рекомендовано
        paragraph_is_recommendation: true

        sentence:
        sentence_type: regular
        """
        counter = 1
        self.list_glob.append("\nparagraph: Рекомендовано\n"
                              "paragraph_is_recommendation: true\n\n"
                              "\tsentence:\n"
                              "\tsentence_type: regular\n\n"
                              "\t\tsubsentence_type: checkbox\n"
                              "\t\tsubsentence_title:\n")
        while "<" not in conclusion.paragraphs[ind + counter].text:
            counter += self.start_end_text(ind, conclusion)
            try:
                part = conclusion.paragraphs[ind + counter].text.split(' [')
                part_1 = part[1].rstrip("] //")
                self.list_glob.append(f"\t\t\tcheckbox_ui_title: {part[0]}\n"
                                      f"\t\t\tcheckbox_selected_text: {part_1}"
                                      )
            except IndexError:
                print(f"Ошибка recommendation в предложении: {conclusion.paragraphs[ind + counter].text}")
            counter += 1

    def conclusion_text(self, ind: int, conclusion: any):
        """
        Генерация текста Заключения
        :param ind: Номер строки в тексте
        :param conclusion: текст шаблона
        :return: None
        """
        counter = 1
        concl = conclusion.paragraphs[ind].text.lstrip("<CONCLUSIONTEXT> ")
        self.list_glob.append(f"\nconclusion_text: {concl}\n")
        while "<" not in conclusion.paragraphs[ind + counter].text:
            try:
                part = conclusion.paragraphs[ind + counter].text.split(' [')
                part_1 = part[1].rstrip("]")
                self.list_glob.append(f"conclusion_conditional: {part[0]}: {part_1}")
            except IndexError:
                print(f"Ошибка conclusion_text в предложении: {conclusion.paragraphs[ind + counter].text}")
            counter += 1

    def save_to_word(self, path: str):
        """
        Сохранение шаблона в директорию
        :param path: путь к директории с исходным документом для шаблонизации
        :return: None
        """
        file = open(path.replace(".docx", ".py"), 'a')
        for text in self.list_glob:
            file.write(text + '\n')
        file.close()


cl = Template()
cl.template()
