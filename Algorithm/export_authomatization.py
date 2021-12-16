"""
Здесь будет алгоритм для автоматической выгрузки шаблонов и заключений в базу данных
"""
from selenium.webdriver import Chrome
from selenium.webdriver.chrome.options import Options
from docx import Document
from time import sleep
import os


def export(link, root, recursion_way):
    opts = Options()
    opts.set_headless()
    assert opts.headless
    file = open(root + '/log.txt', 'a')
    # TODO: передать в функцию название файла и путь к нему (converted_ + '-')
    for dirpath, dirnames, filenames in os.walk(recursion_way):
        for dirname in dirnames:
            if dirname.endswith("-"):
                current_path = os.path.join(dirpath, dirname)
                for doc_name in os.listdir(current_path):
                    try:
                        if ".docx" in doc_name:
                            full_text1 = []
                            full_text2 = []
                            path2 = current_path + '/' + doc_name
                            document = Document(path2)
                            # Добавляем строку без пробелов в начале и конце
                            if len(document.paragraphs[0].text.split(' \n')) == 3:
                                full_text1.append(
                                    document.paragraphs[0].text.split(' \n')[0].strip('\t').strip().strip(
                                        '_') + '//' + dirname + '\n'
                                    + document.paragraphs[0].text.split(' \n')[1].strip('\t').strip().strip('_') + '\n'
                                    + document.paragraphs[0].text.split(' \n')[2].strip('\t').strip().strip('_'))
                                for i in range(1, len(document.paragraphs)):
                                    full_text2.append(
                                        "" + document.paragraphs[i].text.strip().strip('\t').strip('_').replace('\t',
                                                                                                                ' ') + '\n')
                                print(full_text1)
                                # print(full_text2)
                                # print()

                            elif len(document.paragraphs[0].text.split(' \n')) == 2 \
                                    and '-report-text-below-' in document.paragraphs[0].text:
                                full_text1.append(
                                    document.paragraphs[0].text.split('\n')[0].strip().strip('\t').strip(
                                        '_') + '//' + dirname + '\n'
                                    + document.paragraphs[0].text.split('\n')[1].strip().strip('\t').strip('_') + '\n'
                                    + '-report-text-below-')
                                for i in range(1, len(document.paragraphs)):
                                    full_text2.append(
                                        "" + document.paragraphs[i].text.strip().strip('\t').strip('_').replace('\t',
                                                                                                                ' ') + '\n')
                                print(full_text1)
                                # print(full_text2)
                                # print()

                            elif len(document.paragraphs[0].text.split(' \n')) == 2 \
                                    and '-report-text-below-' not in document.paragraphs[0].text:
                                full_text1.append(
                                    document.paragraphs[0].text.split(' \n')[0].strip().strip('\t').strip(
                                        '_') + '//' + dirname + '\n'
                                    + document.paragraphs[0].text.split(' \n')[1].strip('\t').strip().strip('_') + '\n'
                                    + '-report-text-below-')
                                for i in range(2, len(document.paragraphs)):
                                    full_text2.append(
                                        "" + document.paragraphs[i].text.strip().strip('\t').strip('_').replace('\t',
                                                                                                                ' ') + '\n')
                                print(full_text1)
                                # print(full_text2)
                                # print()

                            elif len(document.paragraphs[0].text.split(' \n')) == 1 \
                                    and '-report-text-below-' in document.paragraphs[0].text:
                                full_text1.append(
                                    document.paragraphs[0].text.split('\n')[0].strip().strip('\t').strip(
                                        '_') + '//' + dirname + '\n'
                                    + document.paragraphs[0].text.split('\n')[1].strip().strip('\t').strip('_') + '\n'
                                    + document.paragraphs[0].text.split('\n')[2].strip().strip('\t').strip('_'))
                                for i in range(2, len(document.paragraphs)):
                                    full_text2.append(
                                        "" + document.paragraphs[i].text.strip().strip('\t').strip('_').replace('\t',
                                                                                                                ' ') + '\n')
                                print(full_text1)
                                # print(full_text2)
                                # print()

                            elif len(document.paragraphs[0].text.split(' \n')) == 1 \
                                    and '-report-text-below-' not in document.paragraphs[0].text \
                                    and '-report-text-below-' in document.paragraphs[1].text:
                                full_text1.append(
                                    document.paragraphs[0].text.split('\n')[0].strip().strip('\t') + '//' + dirname + '\n'
                                    + document.paragraphs[0].text.split('\n')[1].strip().strip('\t').strip('_') + '\n'
                                    + '-report-text-below-')
                                for i in range(2, len(document.paragraphs)):
                                    full_text2.append(
                                        "" + document.paragraphs[i].text.strip().strip('\t').strip('_').replace('\t',
                                                                                                                ' ') + '\n')
                                print(full_text1)
                                # print(full_text2)
                                # print()

                            elif len(document.paragraphs[0].text.split(' \n')) == 1 \
                                    and '-report-text-below-' not in document.paragraphs[0].text \
                                    and '-report-text-below-' not in document.paragraphs[1].text:
                                full_text1.append(
                                    document.paragraphs[0].text.split('\n')[0].strip().strip('\t').strip('_') + '//'
                                    + dirname + '\n'
                                    + document.paragraphs[1].text.strip().strip('\t').strip('_') + '\n'
                                    + '-report-text-below-')
                                for i in range(2, len(document.paragraphs)):
                                    full_text2.append(
                                        "" + document.paragraphs[i].text.strip().strip('-report-text-below-\n').strip(
                                            '\t').strip('_').replace('\t',
                                                                     ' ') + '\n')
                                print(full_text1)
                                # print(full_text2)
                                # print()
                            else:
                                print(f'Не попавшее в загрузку: {doc_name}')
                            # browser = Chrome('./chromedriver', options=opts)
                            # browser.get(link)
                            # input_form = browser.find_element_by_name('importtool_text')
                            # input_form.send_keys(full_text1 + list('\n\n') + full_text2)
                            # input_form.submit()
                            # results = browser.find_element_by_xpath("//*[contains(@style,'font-family')]")
                            # file.write(results.text + "\n\n")
                            # browser.close()
                    except ValueError:
                        print(f"ValueError Необработанный документ: {doc_name}")
                    except IndexError:
                        print(f"IndexError Необработанный документ: {doc_name}")
                    else:
                        pass
    file.close()
    quit()
