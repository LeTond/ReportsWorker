from pprint import pprint
from docx import Document
from converting_to_docx import docx_

import os
import requests as re
import easygui


class PersonalReports:
    def __init__(self, url_):
        self.url = url_

    def get_request(self):
        try:
            req = re.get(self.url)
            return req
        except ConnectionError:
            print("Нет соединения с сервером")

    def post_request(self, reports, keys):
        try:
            resp = re.post(
                f"{self.url}"
                f"user_id={keys['user_id']}&"
                f"path={keys['path']}&patient_name={keys['patient_name']}&"
                f"report_date={keys['report_date']}&birthday={keys['birthday']}",
                data=reports)
            return print(f'{resp} Заключение загружено на сервер {keys} \n {reports}')
        except ConnectionError:
            print("Нет соединения с сервером")
        # except re.exceptions.ConnectionError as e:
        #     print("Нет соединения с сервером")
        except ValueError:
            print(f"Ошибка при загрузке данных {keys}")

    @staticmethod
    def import_one_document():
        file = easygui.fileopenbox(msg="Choose a document with report",
                                   default=r"/home/lg/PycharmProjects/DOC_Reading/МРТ/")
        return file

    @staticmethod
    def read_one_file(file_path):
        string = ""
        try:
            document = Document(file_path)
            for par in document.paragraphs:
                string += par.text + '\n'
        except UnicodeDecodeError:
            print("UnicodeDecodeError")
        return string.encode(encoding='utf-8')

    @staticmethod
    def import_documents():
        files = easygui.diropenbox(msg="Choose a directory with reports",
                                   default=r"/home/lg/PycharmProjects/DOC_Reading/")
        return files

    def read_directories(self) -> list:
        file_list = []
        for path, folder, files in os.walk(self.import_documents()):
            for file in files:
                if file.endswith('.docx'):
                    file_list.append([path, file])
        return file_list

    # def read_files(self, file_path):
    #     string = ""
    #     try:
    #         document = Document(file_path)
    #         for par in document.paragraphs:
    #             string += par.text + '\n'
    #     except UnicodeDecodeError:
    #         print("UnicodeDecodeError")
    #     return string.encode(encoding='utf-8')


if __name__ == '__main__':
    keys_ = {'user_id': 10,
             'path': '/МРТ/Импорт',
             'patient_name': 'Test_08',
             'report_date': '2021-09-08',
             'birthday': '2021-01-02'}
    url = "...."
    pr = PersonalReports(url)
    # pprint(pr.post_request(pr.read_one_file(pr.import_one_document()), keys_))
    # pprint(pr.post_request(pr.read_directories(), keys_))
    for j in pr.read_directories():
        pprint(j[1])            ## print report names
