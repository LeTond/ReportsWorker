from selenium.webdriver import Chrome
from selenium.webdriver.chrome.options import Options
from docx import Document


class AutoExport:
    def __init__(self, link, root, current_path, filename, dirname):
        self.link = link
        self.root = root
        self.current_path = current_path
        self.filename = filename
        self.dirname = dirname
        self.full_text1, self.full_text2 = [], []

    def export(self):
        file = open(self.root + '/log.txt', 'a')
        try:
            self.export_data_correction()
            self.send_to_browser(file)
        except ValueError:
            print(f"ValueError Необработанный документ: {self.filename}")
        except IndexError:
            print(f"IndexError Необработанный документ: {self.filename}")
        else:
            pass
        file.close()
        # quit()

    def export_data_correction(self):
        path2 = self.current_path + '/' + self.filename
        document = Document(path2)
        # Добавляем строку без пробелов в начале и конце
        if len(document.paragraphs[0].text.split(' \n')) == 1 \
                and '-report-text-below-' in document.paragraphs[0].text:
            self.full_text1.append(
                document.paragraphs[0].text.split('\n')[0].strip().strip('\t').strip(
                    '_') + '\n'
                + document.paragraphs[0].text.split('\n')[1].strip().strip('\t').strip('_') + '\n'
                + document.paragraphs[0].text.split('\n')[2].strip().strip('\t').strip('_'))
            for i in range(1, len(document.paragraphs)):
                self.full_text2.append(
                    "" + document.paragraphs[i].text.strip().strip('\t').strip('_').replace('\t', ' ') + '\n')
            print(self.full_text1, self.full_text2)
        else:
            print(f'Не попавшее в загрузку: {self.filename}')

    def send_to_browser(self, file):
        opts = Options()
        opts.set_headless()
        assert opts.headless

        browser = Chrome('/home/lg/PycharmProjects/DOC_Reading/chromedriver', options=opts)
        browser.get(self.link)
        input_form = browser.find_element_by_name('importtool_text')
        input_form.send_keys(self.full_text1 + list('\n\n') + self.full_text2)
        input_form.submit()
        results = browser.find_element_by_xpath("//*[contains(@style,'font-family')]")
        file.write(results.text + "\n\n")
        browser.close()
