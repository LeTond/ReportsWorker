# pip install python-docx
from docx import Document
import os
import docx
import shutil


def create_folder_area(anatomy_structure_path: str, root: str):
    """
    Рекурсивное создание директории для Области исследования
    :param anatomy_structure_path: Путь до папки анатомической структуры, н.р.: ХСО - "МРТ//Голова//ХСО"
    :param root: root directory
    :return: None
    """
    new_path = root + '/' + anatomy_structure_path.replace('//', '/')
    try:
        os.makedirs(new_path)
    except FileExistsError:
        print(f"Такая папка существует: {new_path}")
    else:
        pass


def start_segmentation(prepared_document: str, root: str, path: str, key_for_area: list, anatomy_structure_path: str):
    """
    Поиск документов с ключевыми словами
    Проверка наличия достаточных совпадений по ключевым словам и фразам
    Копирование в директорию с ОИ
    :param prepared_document: document path into structure directory
    :param root: root directory
    :param path: Copy directory path
    :param key_for_area: list of key words for anatomical structure
    :param anatomy_structure_path: path to anatomical structure directory
    :return: None
    """
    compare_set_list = []
    control_set = [k for k in range(len(key_for_area))]
    new_path = root + '/' + anatomy_structure_path.replace('//', '/')
    try:
        if ".docx" in prepared_document:
            document_path = root + path + '/' + prepared_document
            conclusion = Document(document_path)
            for para in conclusion.paragraphs:
                for i in range(len(key_for_area)):
                    for j in range(len(key_for_area[i])):
                        if key_for_area[i][j] in para.text or key_for_area[i][j].capitalize() in para.text:
                            compare_set_list.append(i)
                            break
                if set(compare_set_list) == set(control_set):
                    new_current_path = new_path + '/' + '(new)' + prepared_document
                    shutil.copy(document_path, new_current_path)
    except ValueError:
        print(f"Необработанный документ: {prepared_document}")
    else:
        pass


def create_folder_pathology(anatomy_structure_path: str, root: str, key_name: str):
    """
    Создание директорий одноименных ключевым словам в папке патологии
    :param anatomy_structure_path: path to directory
    :param root: root directory
    :param key_name: key word
    :return: None
    """
    new_path = root + '/' + anatomy_structure_path.replace('//', '/') + '/' + key_name
    try:
        os.makedirs(new_path)
        os.mkdir(root + '/' + anatomy_structure_path.replace('//', '/') + '/Прочее')
        os.mkdir(root + '/' + anatomy_structure_path.replace('//', '/') + '/Trash')
    except FileExistsError:
        print("")
    else:
        pass


def continue_segmentation(prepared_document: str, root: str, key_words_for_remove: tuple, pathology_key: str,
                          anatomy_structure_path: str, key_name: str, corn_path: str, conclusion_key_words: list):
    """
    Поиск искомых документов по ключевым словам
    Удаление ненужных предложений
    Сохранение в новый текстовый документ
    Копирование в новую директорию
    :param conclusion_key_words: conclusion key words
    :param prepared_document: document path into pathology directory
    :param root: root directory
    :param key_words_for_remove: key words for remove sentence
    :param pathology_key: key words for pathology
    :param anatomy_structure_path: directory with anatomic structure folders
    :param key_name: key word
    :param corn_path: path to directory with all pathology
    :return: None
    """
    control_set = [k for k in range(len(pathology_key))]
    curr_path = root + '/' + corn_path.replace('//', '/')
    new_path = root + '/' + anatomy_structure_path.replace('//', '/') + '/' + key_name
    counter = len(os.listdir(root + '/' + anatomy_structure_path.replace('//', '/') + '/' + key_name))
    try:
        if ".docx" in prepared_document:
            document_path = curr_path + '/' + prepared_document
            conclusion = Document(document_path)
            number_paragraphs = len(conclusion.paragraphs)

            compare_set_list = determ_key_words_set(conclusion, pathology_key, number_paragraphs, conclusion_key_words)

            new_doc_step1 = docx.Document()
            if set(compare_set_list) == set(control_set):
                if counter < 20 or key_name == 'Прочее':
                    find_sentences_with_key_word(key_words_for_remove, conclusion, number_paragraphs)

                    text = f"{anatomy_structure_path}//{key_name}\n{key_name}\n-report-text-below-"
                    new_doc_step1.add_paragraph(text)
                    create_new_document(new_doc_step1, number_paragraphs, conclusion, new_path, prepared_document)
                elif counter >= 20 and key_name != 'Прочее':
                    new_current_path = root + '/' + anatomy_structure_path.replace('//', '/') + \
                                       '/Trash/' + prepared_document
                    shutil.copy(document_path, new_current_path)

    except ValueError:
        print(f"Необработанный документ: {prepared_document}")
    else:
        pass


def find_sentences_with_key_word(key_words_for_remove, key_head_words, conclusion, number_paragraphs):
    """
    Отфильтровываем всё до параграфа со словом пациент ...
    Отфильтровываем, параграфы со словами врач и тд.....
    :param key_head_words: key words for remove head
    :param key_words_for_remove: key words for remove sentence
    :param conclusion: Document
    :param number_paragraphs: count of paragraphs in Document
    :return: None
    """
    for sentence_number in range(number_paragraphs):
        for key_word in key_head_words:
            if key_word in conclusion.paragraphs[sentence_number].text \
                    or key_word.capitalize() in conclusion.paragraphs[sentence_number].text:
                for index_key_word in range(sentence_number):
                    conclusion.paragraphs[index_key_word].text = None
        for remove_key in key_words_for_remove:
            if remove_key in conclusion.paragraphs[sentence_number].text \
                    or remove_key.capitalize() in conclusion.paragraphs[sentence_number].text:
                conclusion.paragraphs[sentence_number].text = None


def create_new_document(new_doc_step1, number_paragraphs, conclusion, new_path, prepared_document):
    """
    Создаем новый текстовый документ на базе предыдущего без включения пустых параграфов
    :param new_doc_step1: new document
    :param number_paragraphs: count of paragraphs in Document
    :param conclusion: old Document
    :param new_path: new directory
    :param prepared_document: document path into pathology directory
    :return: None
    """
    for sentence_number in range(number_paragraphs):
        if conclusion.paragraphs[sentence_number].text == "":
            pass
        else:
            new_doc_step1.add_paragraph(
                conclusion.paragraphs[sentence_number].text.strip().strip('\t').strip('_').replace('\t', ' '))
        new_doc_step1.save(new_path + '/' + prepared_document)


def determ_key_words_set(conclusion, pathology_key: str, number_paragraphs: int, conclusion_key_words: list) -> list:
    """
    Create set list with words founded in conclusion
    :param conclusion: full conclusion text
    :param pathology_key: key words for pathology
    :param number_paragraphs: number of paragraphs
    :param conclusion_key_words: conclusion key words
    :return: set of key words founded in conclusion
    """
    compare_set_list = []
    for sentence_number in range(number_paragraphs):
        for prob in conclusion_key_words:
            if prob in conclusion.paragraphs[sentence_number].text \
                    or prob.capitalize() in conclusion.paragraphs[sentence_number].text:
                inline_paragraph = sentence_number
                compare_set_list.append(
                    determ_key_word_str(inline_paragraph, conclusion, pathology_key, number_paragraphs))
                break
    return compare_set_list


def determ_key_word_str(inline_paragraph: int, conclusion, pathology_key: str, number_paragraphs: int) -> str:
    """
    Compare key words with conclusion words and return
    :param inline_paragraph: number of paragraph with key word
    :param conclusion: full conclusion text
    :param pathology_key: key words for pathology
    :param number_paragraphs: number of paragraphs
    :return: list with key word from key_list_for_pathology founded in sentence of conclusion
    """
    key_words_str = ""
    for sentence_number in range(inline_paragraph, number_paragraphs):
        for i in range(len(pathology_key)):
            for j in range(len(pathology_key[i])):
                if pathology_key[i][j] in conclusion.paragraphs[sentence_number].text \
                        or pathology_key[i][j].capitalize() in conclusion.paragraphs[sentence_number].text:
                    key_words_str = i
                    break
    return key_words_str


def remove_segmented(prepared_document: str, root: str, anatomy_structure_path: str, key_name: str, corn_path: str):
    """
    Removing documents that contains in new directory or trash directory
    :param prepared_document: Excess documents in anatomical structure directory
    :param root: root directory
    :param anatomy_structure_path: directory with anatomic structure folders
    :param key_name: key word
    :param corn_path: path to directory with all pathology
    :return: None
    """
    curr_path = root + '/' + corn_path.replace('//', '/')
    new_path = root + '/' + anatomy_structure_path.replace('//', '/') + '/' + key_name
    trash = root + '/' + anatomy_structure_path.replace('//', '/') + '/Trash'
    if prepared_document in os.listdir(new_path) or prepared_document in os.listdir(trash):
        os.remove(curr_path + '/' + prepared_document)


def segment_else(prepared_document: str, root: str, key_words_for_remove: tuple, anatomy_structure_path: str):
    """
    Выявляем патологии не попавшие в список с ключевыми словами
    :param prepared_document:
    :param root:
    :param key_words_for_remove:
    :param anatomy_structure_path:
    :return:
    """
    curr_path = root + '/' + anatomy_structure_path.replace('//', '/')
    new_path = root + '/' + anatomy_structure_path.replace('//', '/') + '/Прочее'

    document_path = curr_path + '/' + prepared_document
    try:
        if ".docx" in prepared_document:
            conclusion = Document(document_path)
            new_doc_step1 = docx.Document()
            number_paragraphs = len(conclusion.paragraphs)

            find_sentences_with_key_word(key_words_for_remove, conclusion, number_paragraphs)

            text = f"{anatomy_structure_path} \n Прочее \n -report-text-below-"
            new_doc_step1.add_paragraph(text)

            create_new_document(new_doc_step1, number_paragraphs, conclusion, new_path, prepared_document)
            os.remove(curr_path + '/' + prepared_document)
    except ValueError:
        print(f"Необработанный документ: {prepared_document}")
    else:
        pass
