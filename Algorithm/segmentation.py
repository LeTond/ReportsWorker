from Keys.global_keys import *
from converting_to_docx import docx_
from personal_report import *

from pprint import pprint
import time
from datetime import datetime
from docx import Document

import os
import docx


class DocumentFilter:
    def __init__(self, _root: str, _report_name: list):
        """
        :param _root: root directory
        """
        self.root = _root
        self.path = _report_name[0]
        self.report_name = _report_name[1]
        self.old_document = self.read_document()

    def create_folder_area(self):
        """
        Рекурсивное создание директории для Области исследования
        :return: None
        """
        new_path = self.root + '/' + 'ToExport'
        try:
            os.makedirs(new_path)
        except FileExistsError:
            pass
        else:
            pass

    def read_document(self):
        conclusion = Document(self.path + '/' + self.report_name)
        return conclusion.paragraphs

    def number_paragraphs(self) -> int:
        number_paragraphs = len(self.read_document())
        return number_paragraphs

    def print_document(self):
        for i in range(self.number_paragraphs()):
            pprint(self.read_document()[i].text)

    def filter_header(self):
        """
        Отфильтровываем всё до параграфа со словом пациент ...
        Отфильтровываем, параграфы со словами врач и тд.....
        :return: None
        """
        try:
            self.create_folder_area()
            if ".docx" in self.report_name:
                for paragraph in range(self.number_paragraphs()):
                    for key_word in key_head_words:
                        if key_word in self.old_document[paragraph].text \
                                or key_word.capitalize() in self.old_document[paragraph].text:
                            self.old_document[paragraph].text = None
                            for index_key_word in range(paragraph):
                                self.old_document[index_key_word].text = None
                    for remove_key in key_words_for_remove:
                        if remove_key in self.old_document[paragraph].text \
                                or remove_key.capitalize() in self.old_document[paragraph].text:
                            self.old_document[paragraph].text = None
                self.create_new_document(self.old_document)
        except ValueError:
            print(f"Необработанный документ: {self.read_document()}")
        else:
            pass

    def create_new_document(self, filter_document):
        """
        Создаем новый текстовый документ на базе предыдущего без включения пустых параграфов
        :return: None
        """
        new_document = docx.Document()
        for paragraph in range(len(filter_document)):
            if filter_document[paragraph].text == "":
                pass
            else:
                new_document.add_paragraph(
                    filter_document[paragraph].text.strip('\t').strip('_').replace('\t', ' '))
        new_document.save(self.root + '/ToExport/' + self.report_name)

    def extract_patient_data(self, old_document, key_words):
        data_ = ''
        try:
            for paragraph in range(self.number_paragraphs()):
                for kw in key_words:
                    if kw in old_document[paragraph].text or kw.capitalize() in old_document[paragraph].text:
                        data_ = old_document[paragraph].text.split(':')[1]
                        data_ = ''.join(i for i in data_ if not i.isalpha()) \
                            .lstrip(' ').rstrip('.').rstrip(',').rstrip(' ')
                        pattern_in = "%d.%m.%Y"
                        pattern_out = "%Y-%m-%d"
                        old_date = datetime.strptime(data_, pattern_in)
                        data_ = datetime.strftime(old_date, pattern_out)
                        break
            time.strptime(data_, "%Y-%m-%d")
        except ValueError:
            if len(data_) == 4:
                data_ = f"{data_}-01-01"
            else:
                data_ = "2000-01-01"
        except IndexError:
            print(f'Index error in document {self.report_name}')
        return data_

    def parse_header(self):
        parser_dict = {}
        old_document = self.read_document()
        patient_name = self.report_name.rstrip('.docx')
        parser_dict.update({'report_date': self.extract_patient_data(old_document, key_date_of_study)})
        parser_dict.update({'birthday': self.extract_patient_data(old_document, key_date_of_birth)})
        parser_dict.update({'patient_name': patient_name})
        parser_dict.update({'path': '/МРТ/Импорт'})
        parser_dict.update({'user_id': 80})
        return parser_dict


if __name__ == '__main__':
    # url = "https://kharlinmobile.com:3001/api/import/reports/import-for-user?"
    # docx_(
    #     root,
    #     '/home/lg/PycharmProjects/DOC_Reading/Шмедык НЮ',
    #     '/home/lg/PycharmProjects/DOC_Reading/Шмедык НЮ2'
    # )
    url = "https://client.quickradiology.com:3001/api/import/reports/import-for-user?"
    pr = PersonalReports(url)
    for r_n in pr.read_directories():
        d_f = DocumentFilter(root, r_n)
        keys = d_f.parse_header()
        print(keys)
        d_f.filter_header()
        pr.post_request(pr.read_one_file(root + '/ToExport/' + r_n[1]), keys)
